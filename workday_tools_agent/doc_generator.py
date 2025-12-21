import os
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
from io import BytesIO

import warnings
from docx import Document
from jinja2 import Environment, StrictUndefined
try:
    # Silence docxcompose/pkg_resources deprecation noise during import
    warnings.filterwarnings(
        "ignore",
        category=UserWarning,
        module=r"docxcompose\.properties",
    )
    from docxtpl import DocxTemplate  # Optional, used for templated docs
    _HAS_DOXCTPL = True
except Exception:
    DocxTemplate = None  # type: ignore
    _HAS_DOXCTPL = False

try:
    from docx2pdf import convert as docx2pdf_convert  # Optional, for PDF output on Windows/Word
    _HAS_DOCX2PDF = True
except Exception:
    docx2pdf_convert = None  # type: ignore
    _HAS_DOCX2PDF = False

try:
    from docxtpl import InlineImage  # type: ignore
    from docx.shared import Mm  # type: ignore
    _HAS_INLINE_IMAGE = True
except Exception:
    InlineImage = None  # type: ignore
    Mm = None  # type: ignore
    _HAS_INLINE_IMAGE = False


BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

# In-memory document cache: maps doc_key -> {"bytes": BytesIO, "filename": str}
_document_cache: Dict[str, Dict[str, Any]] = {}


def _sanitize_filename(name: str, preserve_spaces: bool = False) -> str:
    """Sanitize filename by removing only invalid filesystem characters.
    
    Args:
        name: Original filename
        preserve_spaces: If True, keep spaces; otherwise convert to underscores
    
    Returns:
        Sanitized filename
    """
    # Remove only invalid Windows filename characters: < > : " / \ | ? *
    invalid_chars = '<>:"/\\|?*'
    safe = "".join(c for c in name if c not in invalid_chars)
    safe = safe.strip()
    if not preserve_spaces:
        safe = safe.replace(" ", "_")
    return safe or f"document_{int(datetime.now().timestamp())}"


def generate_docx(title: str, content: str, filename: Optional[str] = None) -> Dict[str, str]:
    """Generate a simple DOCX document and return file metadata.

    Args:
        title: Document title (adds a heading)
        content: Main body text content
        filename: Optional desired filename (without path). If not provided, auto-generates.

    Returns:
        Dict with keys: success, filename, download_key
    """
    title = (title or "Document").strip()
    content = (content or "").strip()

    base_name = _sanitize_filename(filename or f"{title.lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    if not base_name.lower().endswith(".docx"):
        base_name += ".docx"

    doc = Document()
    doc.add_heading(title, level=1)
    if content:
        for line in content.splitlines():
            doc.add_paragraph(line)
    
    # Save to BytesIO instead of disk
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Store in cache with unique key
    doc_key = f"{datetime.now().timestamp()}_{base_name}"
    _document_cache[doc_key] = {"bytes": output, "filename": base_name}

    return {
        "success": True,
        "filename": base_name,
        "download_key": doc_key,
    }


def list_templates() -> List[str]:
    """List available .docx templates in the templates folder."""
    return sorted([p.name for p in TEMPLATES_DIR.glob("*.docx")])


def get_template_placeholders(template_name: str) -> List[str]:
    """Return a sorted list of placeholder variables used in the template.

    Tries docxtpl first; if parsing fails, falls back to regex scanning.
    """
    tpl_path = TEMPLATES_DIR / template_name
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {template_name}")

    if _HAS_DOXCTPL:
        try:
            tpl = DocxTemplate(str(tpl_path))
            vars_set = tpl.get_undeclared_template_variables()
            return sorted(list(vars_set))
        except Exception:
            pass

    doc = Document(str(tpl_path))
    return _regex_scan_placeholders(doc)


def generate_docx_from_template(template_name: str, context: Dict[str, Any], filename: Optional[str] = None) -> Dict[str, Any]:
    """Render a DOCX from a .docx template and store in memory.

    Args:
        template_name: Filename of the template inside templates/
        context: Dict of placeholders to values
        filename: Optional output filename

    Returns:
        Dict with keys: success, filename, download_key (for download route)
    """
    if not _HAS_DOXCTPL:
        raise RuntimeError("docxtpl is not installed. Please install 'docxtpl'.")

    tpl_path = TEMPLATES_DIR / template_name
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {template_name}")

    # Use preserve_spaces=True for user-friendly filenames like "Employment Verification Letter - John Doe.docx"
    preserve_spaces = filename is not None and " - " in filename
    base_name = _sanitize_filename(filename or f"render_{Path(template_name).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", preserve_spaces=preserve_spaces)
    if not base_name.lower().endswith(".docx"):
        base_name += ".docx"

    env = Environment(undefined=StrictUndefined, autoescape=False)
    tpl = DocxTemplate(str(tpl_path))

    # Attach images (signature/header) if requested (after tpl is created so InlineImage binds correctly)
    if _HAS_INLINE_IMAGE:
        image_specs = [
            {
                "filenames": ["hr_signature_image", "hr_signature"],
                "width_key": "hr_signature_width_mm",
                "default_width": 40,
            },
            {
                "filenames": ["hr_header_image", "Header", "hr_header"],
                "width_key": "hr_header_width_mm",
                "default_width": 170,
            },
        ]
        for spec in image_specs:
            # First non-empty string value found among the keys
            filename_val = next((context.get(k) for k in spec["filenames"] if isinstance(context.get(k), str) and context.get(k)), None)
            if not filename_val:
                continue
            img_path = TEMPLATES_DIR / filename_val
            if not img_path.exists():
                continue
            try:
                width_mm = context.get(spec["width_key"])
                img_width = Mm(float(width_mm)) if width_mm else Mm(spec["default_width"])
                inline_img = InlineImage(tpl, str(img_path), width=img_width)
                for key in spec["filenames"]:
                    context[key] = inline_img
            except Exception as e:
                print(f"[DocGen] Inline image error for {filename_val}: {e}")

    tpl.render(context or {}, jinja_env=env)
    
    # Save to BytesIO instead of disk
    output = BytesIO()
    tpl.save(output)
    output.seek(0)
    
    # Store in cache with unique key
    doc_key = f"{datetime.now().timestamp()}_{base_name}"
    _document_cache[doc_key] = {"bytes": output, "filename": base_name}

    return {
        "success": True,
        "filename": base_name,
        "download_key": doc_key,
    }


def _maybe_convert_docx_to_pdf(doc_bytes: BytesIO, base_name: str) -> Optional[Dict[str, Any]]:
    """Convert an in-memory DOCX to PDF if docx2pdf is available (Windows + Word).

    Returns a dict matching the cache entry or None if conversion fails.
    """
    if not _HAS_DOCX2PDF:
        return None
    import tempfile
    import shutil

    tmp_dir = tempfile.mkdtemp(prefix="evl_pdf_")
    try:
        docx_path = Path(tmp_dir) / base_name
        pdf_name = Path(base_name).with_suffix(".pdf").name
        pdf_path = Path(tmp_dir) / pdf_name

        with open(docx_path, "wb") as f:
            f.write(doc_bytes.getvalue())

        print(f"[DocGen] Converting DOCX to PDF via Word: {docx_path} -> {pdf_path}")
        docx2pdf_convert(str(docx_path), str(pdf_path))

        if not pdf_path.exists():
            print("[DocGen] PDF conversion failed: output file not found")
            return None

        pdf_bytes = BytesIO(pdf_path.read_bytes())
        pdf_bytes.seek(0)
        print(f"[DocGen] PDF conversion succeeded: {pdf_name}")
        return {"bytes": pdf_bytes, "filename": pdf_name}
    except Exception as e:
        print(f"[DocGen] PDF conversion error: {e}")
        return None
    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


def generate_docx_from_template_as_pdf(template_name: str, context: Dict[str, Any], filename: Optional[str] = None) -> Dict[str, Any]:
    """Render DOCX from template, then convert to PDF if possible. Falls back to DOCX."""
    doc_result = generate_docx_from_template(template_name, context, filename)
    doc_key = doc_result.get("download_key")
    if not doc_key:
        return doc_result

    cache_entry = _document_cache.get(doc_key)
    if not cache_entry:
        return doc_result

    require_pdf = os.getenv("EVL_REQUIRE_PDF", "false").lower() in ("1", "true", "yes")

    pdf_entry = _maybe_convert_docx_to_pdf(cache_entry["bytes"], cache_entry["filename"])
    if not pdf_entry:
        print("[DocGen] PDF conversion unavailable; returning DOCX")
        if require_pdf:
            raise RuntimeError("PDF conversion failed; ensure Microsoft Word is installed and accessible for docx2pdf.")
        doc_result["format"] = "docx"
        doc_result["conversion_failed"] = True
        return doc_result

    pdf_key = f"{datetime.now().timestamp()}_{pdf_entry['filename']}"
    _document_cache[pdf_key] = pdf_entry

    return {
        "success": True,
        "filename": pdf_entry["filename"],
        "download_key": pdf_key,
        "format": "pdf",
        "fallback_docx_key": doc_key,
    }


def get_document_from_cache(doc_key: str) -> Optional[BytesIO]:
    """Retrieve a document from memory cache by key."""
    cache_entry = _document_cache.get(doc_key)
    if cache_entry:
        return cache_entry.get("bytes")
    return None


def get_document_filename_from_cache(doc_key: str) -> Optional[str]:
    """Retrieve the filename of a cached document by key."""
    cache_entry = _document_cache.get(doc_key)
    if cache_entry:
        return cache_entry.get("filename")
    return None


def get_document_mimetype_from_cache(doc_key: str) -> str:
    """Infer mimetype from cached filename."""
    filename = get_document_filename_from_cache(doc_key) or ""
    if filename.lower().endswith(".pdf"):
        return "application/pdf"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def clear_document_cache(doc_key: str):
    """Remove a document from memory cache."""
    _document_cache.pop(doc_key, None)
